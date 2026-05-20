from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 1. 创建Word文档
doc = Document()

# --------------------------
# 全局字体设置（统一宋体）
# --------------------------
def set_font(run, size=12, bold=False):
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# --------------------------
# 标题1：专项资料（居中、大号、加粗）
# --------------------------
def add_title1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=16, bold=True)
    doc.add_paragraph()  # 空行

# --------------------------
# 标题2：大类名称（加粗）
# --------------------------
def add_title2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    doc.add_paragraph()

# --------------------------
# 目录条目（正常文本）
# --------------------------
def add_item(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=12, bold=False)

# ==========================
# 开始生成目录内容
# ==========================

# 政策落实类
add_title1("专项资料")
add_title2("（一）政策落实类")
add_item("1. 农村集体三资管理年度工作总结")
add_item("2. 三资管理工作专项汇报材料及报告")
doc.add_page_break()  # 分页

# 资金管理类
add_title1("专项资料")
add_title2("（二）资金管理类")
add_item("1. 村集体财务管理制度、审批流程、报账制度、村务监督委员会制度")
add_item("2. 会计凭证、记账凭证、银行对账单、现金日记账、银行存款日记账")
add_item("3. 财务报表、收支明细账、债权债务台账")
add_item("4. 大额资金支出审批单、会议记录、合同、付款凭证")
add_item("5. 村级补助资金、项目资金、惠民补贴资金拨付与使用台账")
add_item("6. 集体收入台账（发包收入、租赁收入、补助收入、投资收益等），分年度统计")
add_item("7. 白条入账、坐收坐支、公款私存、违规报销等问题自查材料")
add_item("8. 村务公开（财务公开）资料、公示照片、公开记录")
doc.add_page_break()

# 资产管理类
add_title1("专项资料")
add_title2("（三）资产管理类")
add_item("1. 集体资产台账（房屋、厂房、设备、车辆、办公设施等）")
add_item("2. 资产购置、处置、报废、变卖、出租、出借的审批文件与合同")
add_item("3. 资产评估报告、招投标资料、拍卖记录")
add_item("4. 固定资产盘点表、资产变动记录")
add_item("5. 闲置资产、流失资产、账实不符情况说明")
add_item("6. 村级集体经济组织登记证书、营业执照等")
doc.add_page_break()

# 资源管理类
add_title1("专项资料")
add_title2("（四）资源管理类")
add_item("1. 集体土地、林地、水库、荒地等资源台账")
add_item("2. 土地承包、流转、租赁、征收、征用合同及备案资料")
add_item("3. 高标准农田、集体建设用地、宅基地管理资料")
add_item("4. 资源发包招投标、竞价、公示记录")
add_item("5. 征地补偿款分配方案、发放明细、签收记录")
add_item("6. 自然资源权属证明、确权登记资料")
doc.add_page_break()

# 合同管理类
add_title1("专项资料")
add_title2("（五）合同管理类")
add_item("1. 集体所有经济合同（承包、租赁、工程、买卖、服务等）")
add_item("2. 合同台账、履行情况、收款记录")
add_item("3. 无效合同、超期合同、低价发包合同、未备案合同清单")
doc.add_page_break()

# 民主决策与程序类
add_title1("专项资料")
add_title2("（六）民主决策与程序类")
add_item("1. 村党委（总支）会、村“两委”会、党员大会、村民代表会议记录本")
add_item("2. “四议两公开”记录、签到表、表决票、决议文件、公开资料等")
add_item("3. 重大事项决策台账（工程建设、资产处置、大额支出等）")
add_item("4. 村务监督委员会履职记录、审核意见")
doc.add_page_break()

# 工程项目类
add_title1("专项资料")
add_title2("（七）工程项目类")
add_item("1. 村级工程项目立项、审批、招投标、合同、验收资料")
add_item("2. 工程预决算、审计报告、付款凭证")
add_item("3. 工程发包、分包、变更签证资料")
add_item("4. 项目资金使用明细、是否存在虚报冒领")
doc.add_page_break()

# 三资监管平台类
add_title1("专项资料")
add_title2("（八）“三资”监管平台、报账系统资料")
add_item("1. 农村集体“三资”监管系统录入数据、截图")
add_item("2. 线上审批、线上公开记录")
add_item("3. 数据与账实一致性自查材料")
doc.add_page_break()

# 党风廉政建设类
add_title1("专项资料")
add_title2("（九）党风廉政建设类")
add_item("1. 推进惩治和预防腐败体系建设、履行党风廉政建设及正风肃纪、执纪问责等情况")
add_item("2. 专项监督检查和专项治理情况，如，开展群众身边的腐败问题和不正之风问题监督检查情况及相关材料")
doc.add_page_break()

# 信访与问题整改类
add_title1("专项资料")
add_title2("（十）信访与问题整改类")
add_item("1. 以往巡察、审计部门（含上级专项审计）出具的巡察、审计报告及整改报告和佐证资料")
add_item("2. 信访举报、矛盾纠纷、群众反映问题及办理结果台账，工作记录本")
add_item("3. 以往各类督查、检查反馈问题及整改报告和佐证材料")

# 保存文件
doc.save("农村三资巡察专项资料目录.docx")
print("Word文件生成完成：农村三资巡察专项资料目录.docx")